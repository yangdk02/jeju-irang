from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SOURCE = Path(r"G:\내 드라이브\Work\Career Assets\Awards & Projects\생성형 AI 기반 바이브코딩 활용 앱 서비스 제작 입문 과정\02_worksheets\[워크시트] 3일차 화면이 된 데이터_전강양2.docx")
OUTPUT = Path(r"C:\Users\yangd\projects\jeju-irang\outputs\[워크시트] 3일차 화면이 된 데이터_전강양2_작성완료.docx")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_paragraph(paragraph, value: str):
    donor_rpr = None
    for run in paragraph.runs:
        if run._r.rPr is not None:
            donor_rpr = deepcopy(run._r.rPr)
            break
    clear_paragraph(paragraph)
    run = paragraph.add_run(value)
    if donor_rpr is not None:
        run._r.insert(0, donor_rpr)


def set_cell(cell, value: str):
    paragraph = cell.paragraphs[0]
    set_paragraph(paragraph, value)
    for extra in cell.paragraphs[1:]:
        clear_paragraph(extra)


def ensure_data_rows(table, count: int):
    while len(table.rows) - 1 < count:
        template = deepcopy(table.rows[-1]._tr)
        table._tbl.append(template)
        row = table.rows[-1]
        for cell in row.cells:
            set_cell(cell, "")


doc = Document(SOURCE)

# 1. Data summary: retain the worksheet's convention while clarifying header count.
summary = doc.tables[4]
set_cell(summary.cell(0, 3), "24개 데이터 행 (헤더 포함 25줄)")
set_cell(summary.cell(1, 1), "29")
set_cell(summary.cell(1, 3), "minimum_age (13개), review_summary (4개)")

# 2. Explain every CSV column, including columns currently loaded but not displayed.
column_uses = [
    ("place_id", "목록 선택값·즐겨찾기 연결에 쓰는 내부 ID"),
    ("place_name", "장소명 검색, 목록 제목, 상세 제목"),
    ("category_level_2", "사이드바 시설 유형 필터, 목록·상세 배지"),
    ("city_name", "시·군·구 위치 정보 (현재 상세 화면에는 미표시)"),
    ("legal_dong_name", "상세 카드의 읍·면·동 위치 표시"),
    ("region_group", "상단 지역 선택 필터, 목록·상세 지역 표시"),
    ("road_address", "상세 카드의 주소"),
    ("latitude", "지도용 위도 데이터로 변환 (현재 화면에는 미표시)"),
    ("longitude", "지도용 경도 데이터로 변환 (현재 화면에는 미표시)"),
    ("phone", "상세 카드의 연락처"),
    ("website_url", "상세 카드의 홈페이지 링크"),
    ("closed_days", "상세 카드의 휴무일 지표"),
    ("opening_hours", "목록의 운영시간, 상세 운영시간 지표"),
    ("free_parking", "무료주차 체크 필터, 상세 편의시설 배지"),
    ("paid_parking", "유료주차 여부 데이터 (현재 화면에는 미표시)"),
    ("has_admission_fee", "무료 입장 요약 지표, 이용요금 표시"),
    ("admission_fee", "이용요금 낮은 순 정렬, 목록·상세 최저요금"),
    ("admission_fee_detail", "상세 카드의 연령별 요금 안내"),
    ("has_age_limit", "상세 카드의 연령 제한 여부"),
    ("minimum_age", "상세 카드의 최소 이용 가능 연령"),
    ("nursing_room", "수유실 체크 필터, 상세 편의시설 배지"),
    ("stroller_rental", "유모차 대여 체크 필터, 상세 편의시설 배지"),
    ("reservation_url", "상세 카드의 예약 페이지 링크"),
    ("space_type", "실내·실외 필터, 목록·상세 공간 유형"),
    ("resident_discount", "도민 할인 체크 필터, 상세 편의시설 배지"),
    ("diaper_changing_table", "기저귀교환대 체크 필터, 상세 편의시설 배지"),
    ("photo_url", "상세 카드의 장소 사진"),
    ("description", "검색 대상, 상세 카드의 한 줄 설명"),
    ("review_summary", "검색 대상, 상세 카드의 방문 참고·후기 요약"),
]
column_table = doc.tables[5]
ensure_data_rows(column_table, len(column_uses))
for row_index, (column, usage) in enumerate(column_uses, start=1):
    set_cell(column_table.cell(row_index, 0), column)
    set_cell(column_table.cell(row_index, 1), usage)

# 3. Connect visible UI elements to the data columns they consume.
screen_table = doc.tables[8]
screen_rows = [
    ("검색창", "장소명·설명·후기에서 입력한 키워드를 찾습니다.", "place_name, description, review_summary"),
    ("필터 1", "전체 또는 6개 권역 중 하나를 선택해 장소를 좁힙니다.", "region_group"),
    ("필터 2", "시설 유형·실내외·무료주차·수유실·유모차·기저귀교환대·도민 할인을 선택합니다.", "category_level_2, space_type, free_parking, nursing_room, stroller_rental, diaper_changing_table, resident_discount"),
    ("요약 지표", "검색 결과 수, 무료 입장 수, 실내 장소 수를 보여 줍니다.", "place_id, has_admission_fee, space_type"),
    ("목록 / 표", "장소명·지역·시설 유형·공간·운영시간·이용요금을 비교합니다.", "place_name, region_group, category_level_2, space_type, opening_hours, has_admission_fee, admission_fee"),
    ("상세 카드", "사진, 위치, 설명, 편의시설, 운영시간, 요금, 주소, 연락처, 홈페이지·예약 링크, 후기 등을 보여 줍니다.", "photo_url, place_name, region_group, legal_dong_name, description, free_parking, nursing_room, stroller_rental, diaper_changing_table, resident_discount, opening_hours, closed_days, has_admission_fee, admission_fee, admission_fee_detail, has_age_limit, minimum_age, road_address, phone, website_url, reservation_url, review_summary"),
    ("차트 (있다면)", "현재 화면에는 차트를 사용하지 않습니다.", "없음"),
]
for row_index, values in enumerate(screen_rows, start=1):
    for col_index, value in enumerate(values):
        set_cell(screen_table.cell(row_index, col_index), value)

# 4. Day 3 completed feature and remaining MVP work.
feature_table = doc.tables[9]
set_cell(feature_table.cell(0, 1), "아이 동반 장소 목록 조회")
set_cell(feature_table.cell(1, 1), "사용자는 CSV의 24개 장소를 목록으로 보고, 검색 결과와 기본 정보를 비교할 수 있습니다.")
set_cell(feature_table.cell(2, 1), "① 지역·시설유형·실내외·육아 편의조건 필터링\n② 장소 상세정보 보기")

# 5. Optional-column removal test. The app uses row.get(), so the photo fallback remains safe.
deletion_table = doc.tables[12]
set_cell(deletion_table.cell(0, 1), "photo_url")
set_cell(deletion_table.cell(1, 1), "실제 장소 사진이 사라지고 '장소 사진 준비 중' 대체 영역이 표시되었습니다.")
set_cell(deletion_table.cell(2, 1), "☑  예                    ☐  아니오")

# Record one real setup error encountered during development and its remedy.
set_paragraph(
    doc.paragraphs[22],
    "해당 없음. 컬럼 삭제 시험에서는 앱이 멈추지 않았습니다.\n오류 기록: 실행 환경에 Streamlit이 없어 ModuleNotFoundError가 발생했습니다. requirements.txt의 streamlit과 pandas를 설치하도록 안내해 해결했습니다.",
)

# 6. Complete all non-screenshot checklist items truthfully.
set_paragraph(doc.paragraphs[26], '☑  컬럼마다 "화면의 어디에 쓰이는지"를 적었습니다')
set_paragraph(doc.paragraphs[27], '☑  화면 요소마다 "어떤 컬럼을 쓰는지"를 적었습니다')
set_paragraph(doc.paragraphs[28], '☐  화면 어디에도 안 쓰이는 컬럼이 없습니다 (city_name, latitude, longitude, paid_parking은 현재 화면 미표시)')
set_paragraph(doc.paragraphs[29], "☑  컬럼을 지워도 앱이 멈추지 않습니다")
set_paragraph(doc.paragraphs[30], "☑  오류 1개를 기록했습니다")
set_paragraph(doc.paragraphs[31], "☑  4일차에 추가할 입력 기능이 정해졌습니다")

# 7. Day 4 input feature.
day4 = doc.tables[15]
set_cell(day4.cell(0, 1), "사용자가 닉네임을 입력하고 선택한 장소를 즐겨찾기에 저장하게 합니다. 저장된 장소는 닉네임별 목록에서 다시 조회할 수 있습니다.")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)

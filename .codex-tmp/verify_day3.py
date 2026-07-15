from pathlib import Path
from zipfile import ZipFile

from docx import Document

source_path = Path(r"G:\내 드라이브\Work\Career Assets\Awards & Projects\생성형 AI 기반 바이브코딩 활용 앱 서비스 제작 입문 과정\02_worksheets\[워크시트] 3일차 화면이 된 데이터_전강양2.docx")
output_path = Path(r"C:\Users\yangd\projects\jeju-irang\outputs\[워크시트] 3일차 화면이 된 데이터_전강양2_작성완료.docx")

source = Document(source_path)
output = Document(output_path)

assert len(output.tables) == 16
assert len(output.paragraphs) == 34
assert len(output.tables[5].rows) == 30
assert output.tables[4].cell(0, 3).text == "24개 데이터 행 (헤더 포함 25줄)"
assert output.tables[4].cell(1, 1).text == "29"
assert output.tables[5].cell(1, 0).text == "place_id"
assert output.tables[5].cell(29, 0).text == "review_summary"
assert output.tables[8].cell(1, 1).text
assert output.tables[9].cell(0, 1).text == "아이 동반 장소 목록 조회"
assert output.tables[12].cell(0, 1).text == "photo_url"
assert "닉네임" in output.tables[15].cell(0, 1).text

# Capture placeholders must remain unchanged.
for table_index in (3, 7, 11):
    assert output.tables[table_index].cell(0, 0).text == source.tables[table_index].cell(0, 0).text

# A successful ZIP scan catches malformed package relationships/parts early.
with ZipFile(output_path) as archive:
    bad = archive.testzip()
    assert bad is None, bad
    assert "word/document.xml" in archive.namelist()

print(f"verified: {output_path.name}")
print(f"size={output_path.stat().st_size} tables={len(output.tables)} column_mapping_rows={len(output.tables[5].rows)-1}")
print("capture_placeholders=unchanged zip=ok reopen=ok")

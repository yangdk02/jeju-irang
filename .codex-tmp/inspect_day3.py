from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree

SOURCE = Path(r"G:\내 드라이브\Work\Career Assets\Awards & Projects\생성형 AI 기반 바이브코딩 활용 앱 서비스 제작 입문 과정\02_worksheets\[워크시트] 3일차 화면이 된 데이터_전강양2.docx")

doc = Document(SOURCE)
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)} inline_shapes={len(doc.inline_shapes)}")

print("\n=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs, 1):
    print(f"P{i:03d} style={p.style.name!r} text={p.text!r}")

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables, 1):
    print(f"\nTABLE {ti}: {len(table.rows)}x{len(table.columns)} style={table.style.name if table.style else None!r}")
    for ri, row in enumerate(table.rows, 1):
        cells = [cell.text.replace("\n", " / ") for cell in row.cells]
        print(f"R{ri:02d}: " + " | ".join(repr(value) for value in cells))

print("\n=== XML TEXT OUTSIDE BODY PARAGRAPH API ===")
with ZipFile(SOURCE) as archive:
    xml = archive.read("word/document.xml")
root = etree.fromstring(xml)
ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
for i, node in enumerate(root.xpath(".//w:sdt | .//w:txbxContent", namespaces=ns), 1):
    text = "".join(node.xpath(".//w:t/text()", namespaces=ns))
    print(f"NODE{i:03d} tag={etree.QName(node).localname} text={text!r}")

print("\n=== BODY ORDER ===")
table_lookup = {table._tbl: i for i, table in enumerate(doc.tables, 1)}
paragraph_lookup = {paragraph._p: i for i, paragraph in enumerate(doc.paragraphs, 1)}
for child in doc.element.body.iterchildren():
    name = etree.QName(child).localname
    if name == "p":
        i = paragraph_lookup.get(child)
        text = "".join(child.xpath(".//w:t/text()"))
        print(f"P{i:03d}: {text!r}")
    elif name == "tbl":
        print(f"TABLE {table_lookup.get(child)}")

from pathlib import Path
from docx import Document

source = Path(r"C:\Users\yangd\Downloads\[워크시트] 2일차 앱 설계서_전강양2.docx")
doc = Document(source)

print("=== PARAGRAPHS ===")
for index, paragraph in enumerate(doc.paragraphs, start=1):
    text = paragraph.text.strip()
    if text:
        print(f"P{index:03d} [{paragraph.style.name}] {text}")

print("\n=== TABLES ===")
for table_index, table in enumerate(doc.tables, start=1):
    print(f"\nTABLE {table_index}: {len(table.rows)} rows x {len(table.columns)} cols")
    for row_index, row in enumerate(table.rows, start=1):
        cells = [cell.text.replace("\n", " / ").strip() for cell in row.cells]
        print(f"R{row_index:02d}: " + " | ".join(cells))

print("\n=== SHAPES ===")
print(f"inline_shapes={len(doc.inline_shapes)}")
